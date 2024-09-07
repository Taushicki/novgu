import os
import zipfile

from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx2pdf import convert
from PyPDF2 import PdfReader
from fastapi import HTTPException
from fastapi.responses import FileResponse


class ZIP:
    @staticmethod
    def create_zip(files: list):
        if len(files) > 1:
            zip_file_name = "output.zip"
            with zipfile.ZipFile(zip_file_name, "w") as zipf:
                for file in files:
                    zipf.write(file)
                    os.remove(file)
                file_path = os.path.join(os.getcwd(), zip_file_name)
                if os.path.exists(file_path):
                    return FileResponse(
                        path=file_path,
                        filename=zip_file_name,
                        media_type="application/zip",
                        headers={
                            "Content-Disposition": f"attachment; filename={zip_file_name}"
                        },
                    )
                else:
                    raise HTTPException(status_code=404, detail="File not found")
        else:
            file_path = os.path.join(os.getcwd(), files[0])
            if os.path.exists(file_path):
                return FileResponse(
                    path=file_path,
                    filename=files[0],
                    media_type="application/pdf",
                    headers={"Content-Disposition": f"attachment; filename={files[0]}"},
                )
            else:
                raise HTTPException(status_code=404, detail="File not found")


class PDF:
    def create_certificate(
        self,
        output_file_name,
        course: str,
        group: str,
        forms: str,
        fio: str,
        fio_upper: str,
        spec: str,
        marks: list[tuple],
        debt: list[str] = None,
    ) -> None:
        # Словарь для замены слов
        replacing = {
            "date": datetime.now().strftime("%d.%m.%Y"),
            "course": course,
            "group": group,
            "forms": forms,
            "fio": fio,
            "fio_upper": fio_upper.upper(),
            "spec": spec,
            "result": "На данный момент академических задолженностей не имеет.",
            "table_list_count": "1",
        }

        # Добавление задолженностей
        if debt:
            replacing["result"] = (
                "На данный момент имеет академические задолженности по дисциплинам:"
            )
            for subject in debt:
                replacing["result"] += f" {subject},"
            replacing["result"] = replacing["result"][:-1]

        # Открытие документа
        doc = Document("base.docx")

        # Добавление заголовков для таблицы
        marks.insert(0, ("Семестр", "Дисциплина", "", "Рейтинг", "Оценка"))
        # Добавление таблицы с дисциплинами
        table = doc.add_table(rows=len(marks), cols=5)
        # Установка ширины колонок
        table.columns[0].width = Inches(0.75)
        table.columns[1].width = Inches(10)
        table.columns[2].width = Inches(0.5)
        table.columns[3].width = Inches(0.625)
        table.columns[4].width = Inches(0.625)
        for column in table.columns:
            for cell in column.cells:
                cell.width = column.width
        # Заполнение таблицы данными и настройка шрифта
        for row_idx, row_data in enumerate(marks):
            for col_idx, item in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                cell.text = item
                # Установка шрифта для ячейки
                for paragraph in cell.paragraphs:
                    run = paragraph.runs[0]
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
        # Установка стиля таблицы
        table.style = "Table Grid"
        # Определение кол-ва листов приложения
        doc.save("updated_base.docx")
        convert("updated_base.docx", "temp.pdf")
        with open("temp.pdf", "rb") as f:
            reader = PdfReader(f)
            replacing["table_list_count"] = len(reader.pages) - 1
        os.remove("temp.pdf")

        # Замена слов
        paragraphs = list(doc.paragraphs)
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraphs.append(paragraph)
        for p in paragraphs:
            for key, val in replacing.items():
                key_name = "${{{}}}".format(key)
                if key_name in p.text:
                    inline = p.runs
                    started = False
                    key_index = 0
                    found_runs = list()
                    found_all = False
                    replace_done = False
                    for i in range(len(inline)):

                        if key_name in inline[i].text and not started:
                            found_runs.append(
                                (i, inline[i].text.find(key_name), len(key_name))
                            )
                            text = inline[i].text.replace(key_name, str(val))
                            inline[i].text = text
                            replace_done = True
                            found_all = True
                            break

                        if key_name[key_index] not in inline[i].text and not started:
                            continue

                        if (
                            key_name[key_index] in inline[i].text
                            and inline[i].text[-1] in key_name
                            and not started
                        ):
                            start_index = inline[i].text.find(key_name[key_index])
                            check_length = len(inline[i].text)
                            for text_index in range(start_index, check_length):
                                if inline[i].text[text_index] != key_name[key_index]:
                                    break
                            if key_index == 0:
                                started = True
                            chars_found = check_length - start_index
                            key_index += chars_found
                            found_runs.append((i, start_index, chars_found))
                            if key_index != len(key_name):
                                continue
                            else:
                                found_all = True
                                break

                        if (
                            key_name[key_index] in inline[i].text
                            and started
                            and not found_all
                        ):
                            chars_found = 0
                            check_length = len(inline[i].text)
                            for text_index in range(0, check_length):
                                if inline[i].text[text_index] == key_name[key_index]:
                                    key_index += 1
                                    chars_found += 1
                                else:
                                    break
                            found_runs.append((i, 0, chars_found))
                            if key_index == len(key_name):
                                found_all = True
                                break

                    if found_all and not replace_done:
                        for i, item in enumerate(found_runs):
                            index, start, length = [t for t in item]
                            if i == 0:
                                text = inline[index].text.replace(
                                    inline[index].text[start : start + length], str(val)
                                )
                                inline[index].text = text
                            else:
                                text = inline[index].text.replace(
                                    inline[index].text[start : start + length], ""
                                )
                                inline[index].text = text

        # Сохранение изменений
        doc.save("updated_base.docx")
        convert("updated_base.docx", output_file_name)
        os.remove("updated_base.docx")
